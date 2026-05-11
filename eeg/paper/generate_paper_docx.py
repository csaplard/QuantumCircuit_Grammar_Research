"""
EEG Grammar Fingerprinting paper -> DOCX (Zenodo-ready).
Style follows the Fisher Path Speed preprint (Csaplar 2026e) exactly:
- Title: large bold, centered
- "Continuation of ..." line: small italic, centered
- Author block (name, affiliation, ORCID, date): centered, small
- Abstract: bold "Abstract" header (smaller than section heads), body justified
- Section heads "1. Introduction": bold, ~13pt
- Subsection heads "2.1 Dataset.": bold inline body-size (NOT heading style)
- Body: 11pt Calibri
- Figure captions: bold "Figure N: title." + plain body
- Table captions: bold "Table N: ..." above table
- References: bullet list
- End: "— End of preprint —" italic gray centered
"""

import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
MD_FILE = os.path.join(HERE, "EEG_Grammar_Fingerprinting.md")
FIG_DIR = os.path.join(HERE, "figures")
OUT = os.path.join(HERE, "EEG_Grammar_Fingerprinting.docx")

FIGURE_FILES = {
    1: "fig2_main_effect.png",
    2: "fig3_topography.png",
    3: "fig5_fingerprints.png",
    4: "fig4_loso.png",
    5: "fig6_aaft.png",
}

FIGURE_CAPTIONS = {
    1: ("Figure 1: Main effect.",
        "(A) Box plots of per-stream perplexity for the 684 (subject, channel) "
        "streams in each regime. Median perplexity is 4.01 at rest and 4.63 during "
        "arithmetic (means 4.05 and 4.69, respectively). (B) Per-stream pair scatter "
        "plot. Each point is one (subject, channel) pair; 87.3% lie above the y = x "
        "diagonal."),
    2: ("Figure 2: Topography.",
        "(A) Topographic delta map (arithmetic - rest) interpolated across the 10-20 "
        "montage. All 19 electrodes show positive difference; the maximum spans the "
        "posterior-parietal-midline ring, the minimum is at frontopolar sites. "
        "(B) Per-channel mean delta with SEM, sorted."),
    3: ("Figure 3: Average fingerprints.",
        "(A) Average 7x7 grammar fingerprint matrix during eyes-closed rest. Rows: "
        "current SAX symbol; columns: predicted next SAX symbol; cell value: "
        "P(next | current). (B) Average fingerprint during mental arithmetic. "
        "(C) Difference matrix (arithmetic - rest); the dominant change is reduced "
        "self-persistence at the central symbol P(d -> d)."),
    4: ("Figure 4: Generalization.",
        "(A) Distribution of leave-one-subject-out classification accuracies across "
        "the 36 folds. Mean = 75.2% +/- 10.6% (chance = 50%). (B) Pooled confusion "
        "matrix across all folds, row-normalized."),
    5: ("Figure 5: AAFT surrogate test.",
        "(A) Boxplots of per-stream AAFT - real perplexity delta for rest and "
        "arithmetic. (B) Real vs AAFT perplexity scatter by regime; the cluster sits "
        "on the y = x diagonal in both regimes. (C) Per-channel mean AAFT - real "
        "delta by regime."),
}

INLINE_BOLD = re.compile(r"\*\*(.+?)\*\*")


def set_paragraph_spacing(p, before=0, after=4):
    """Tighter spacing similar to preprint look."""
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)


def render_inline(paragraph, text, base_size=11):
    """Render text into a paragraph, splitting on **bold** markers."""
    pos = 0
    for m in INLINE_BOLD.finditer(text):
        if m.start() > pos:
            r = paragraph.add_run(text[pos:m.start()])
            r.font.size = Pt(base_size)
        r = paragraph.add_run(m.group(1))
        r.bold = True
        r.font.size = Pt(base_size)
        pos = m.end()
    if pos < len(text):
        r = paragraph.add_run(text[pos:])
        r.font.size = Pt(base_size)


def add_section_heading(doc, text, level):
    """
    Section heading sizing match Fisher Path Speed style:
      level=1: "1. Introduction" - bold 12.5pt
      level=2: "2.1 Dataset." - bold 11pt (inline body-size, not large heading)
    Custom rendering instead of Heading 1/2 styles so we get exact look.
    """
    sizes = {0: 16, 1: 12.5, 2: 11}
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=10 if level == 1 else 8,
                          after=4 if level == 2 else 6)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(sizes.get(level, 11))
    return p


def add_centered(doc, text, size=10, italic=False, bold=False, color=None,
                 space_after=2):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=space_after)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.italic = italic
    r.bold = bold
    if color is not None:
        r.font.color.rgb = color
    return p


def add_table_from_md(doc, header_row, rows, caption=None):
    if caption:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=8, after=4)
        r = p.add_run(caption)
        r.bold = True
        r.font.size = Pt(10)

    table = doc.add_table(rows=len(rows) + 1, cols=len(header_row))
    table.style = 'Light Grid Accent 1'
    for j, h in enumerate(header_row):
        cell = table.rows[0].cells[j]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = ''
            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(10)
    doc.add_paragraph()


def insert_figure(doc, fig_num):
    if fig_num not in FIGURE_FILES:
        return
    fpath = os.path.join(FIG_DIR, FIGURE_FILES[fig_num])
    if not os.path.exists(fpath):
        doc.add_paragraph(f"[FIGURE {fig_num} NOT FOUND]")
        return
    pimg = doc.add_paragraph()
    pimg.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(pimg, before=8, after=4)
    pimg.add_run().add_picture(fpath, width=Inches(6.0))

    title, body = FIGURE_CAPTIONS[fig_num]
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(cap, before=0, after=8)
    r1 = cap.add_run(title + " ")
    r1.font.size = Pt(9)
    r1.font.bold = True
    r2 = cap.add_run(body)
    r2.font.size = Pt(9)


def is_table_row(line):
    return line.strip().startswith('|') and line.strip().endswith('|')


def parse_md_table(lines, start_idx):
    rows = []
    i = start_idx
    while i < len(lines) and is_table_row(lines[i]):
        cells = [c.strip() for c in lines[i].strip().strip('|').split('|')]
        rows.append(cells)
        i += 1
    if len(rows) >= 2 and all(re.match(r'^[-: ]+$', c) for c in rows[1]):
        header = rows[0]
        body = rows[2:]
    else:
        header = rows[0] if rows else []
        body = rows[1:] if len(rows) > 1 else []
    return header, body, i


def add_body_paragraph(doc, text, justify=True):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=6)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    render_inline(p, text, base_size=11)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    set_paragraph_spacing(p, before=0, after=2)
    render_inline(p, text, base_size=10)
    return p


def main():
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # 1.0 line spacing globally
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    with open(MD_FILE, encoding='utf-8') as f:
        lines = f.readlines()

    i = 0
    in_code = False
    code_buffer = []
    # State: pending caption for next table
    pending_table_caption = None

    while i < len(lines):
        line = lines[i].rstrip('\n')
        stripped = line.strip()

        if stripped.startswith('```'):
            if in_code:
                p = doc.add_paragraph()
                set_paragraph_spacing(p, before=4, after=4)
                r = p.add_run('\n'.join(code_buffer))
                r.font.name = 'Consolas'
                r.font.size = Pt(9)
                code_buffer = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_buffer.append(line)
            i += 1
            continue

        # Tables
        if is_table_row(line):
            header, body, i = parse_md_table(lines, i)
            add_table_from_md(doc, header, body, caption=pending_table_caption)
            pending_table_caption = None
            continue

        # Inline figure marker
        m_fig = re.match(r"^\*\*Figure (\d+):", stripped)
        if m_fig:
            fig_num = int(m_fig.group(1))
            insert_figure(doc, fig_num)
            i += 1
            continue

        # Inline table caption like "**Table 1: ...**"
        m_tbl = re.match(r"^\*\*Table (\d+): (.+?)\*\*$", stripped)
        if m_tbl:
            pending_table_caption = f"Table {m_tbl.group(1)}: {m_tbl.group(2)}"
            i += 1
            continue

        # Headings
        if stripped.startswith('# '):
            txt = stripped[2:].strip()
            add_centered(doc, txt, size=16, bold=True, space_after=8)
        elif stripped.startswith('## '):
            txt = stripped[3:].strip()
            # "Abstract" is special - centered/bold smaller
            if txt.lower() == 'abstract':
                p = doc.add_paragraph()
                set_paragraph_spacing(p, before=10, after=4)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run('Abstract')
                r.bold = True
                r.font.size = Pt(11)
            else:
                add_section_heading(doc, txt, level=1)
        elif stripped.startswith('### '):
            txt = stripped[4:].strip()
            add_section_heading(doc, txt, level=2)
        elif stripped.startswith('- '):
            add_bullet(doc, stripped[2:])
        elif not stripped:
            pass  # blank
        else:
            # Special author block lines
            if stripped.startswith('Continuation of '):
                add_centered(doc, stripped, size=9, italic=True, space_after=8)
            elif stripped == 'Daniel Csaplár':
                add_centered(doc, stripped, size=12, space_after=2)
            elif stripped.startswith('Independent Researcher'):
                add_centered(doc, stripped, size=10, space_after=1)
            elif stripped.startswith('ORCID'):
                add_centered(doc, stripped, size=10, space_after=1)
            elif re.match(r'^[A-Z][a-z]+ 20\d{2}$', stripped):
                add_centered(doc, stripped, size=10, space_after=8)
            elif stripped == '— End of preprint —':
                add_centered(doc, stripped, size=10, italic=True,
                              color=RGBColor(0x80, 0x80, 0x80), space_after=4)
            else:
                add_body_paragraph(doc, stripped, justify=True)

        i += 1

    doc.save(OUT)
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    main()
