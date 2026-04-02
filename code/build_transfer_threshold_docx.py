from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.shared import Inches, Pt

REPO_ROOT = Path(__file__).resolve().parents[1]
RESULTS = REPO_ROOT / "results"


def add_preformatted(doc: Document, text: str) -> None:
    for line in text.splitlines():
        p = doc.add_paragraph()
        run = p.add_run(line if line else " ")
        run.font.name = "Courier New"
        run.font.size = Pt(8)


def build_docx() -> Path:
    doc = Document()
    doc.add_heading("Threshold Transfer Model — Paradigm Summary", level=0)
    doc.add_paragraph(
        "Log-linear transfer model: log(N*) = log(H0) + alpha*log(Q) + log(B_backend) + log(C_regime), "
        "fit jointly on SycamoreRandom (Sycamore) and IBM Fisher thresholds (Marrakesh, Torino)."
    )

    report_path = RESULTS / "threshold_transfer_model_report.txt"
    if report_path.exists():
        doc.add_heading("Fit report (verbatim)", level=1)
        add_preformatted(doc, report_path.read_text(encoding="utf-8", errors="replace"))
    else:
        doc.add_paragraph("(threshold_transfer_model_report.txt not found)")

    pred_path = RESULTS / "threshold_transfer_predictions.csv"
    if pred_path.exists():
        df = pd.read_csv(pred_path)
        # De-duplicate identical rows from report pipeline if any
        df = df.drop_duplicates(
            subset=["source", "backend_group", "regime_family", "Q", "N_star"], keep="first"
        )
        doc.add_heading("Predictions vs observed N*", level=1)
        cols = ["source", "backend_group", "regime_family", "Q", "N_star", "N_pred"]
        sub = df[[c for c in cols if c in df.columns]]
        table = doc.add_table(rows=1, cols=len(sub.columns))
        hdr = table.rows[0].cells
        for j, c in enumerate(sub.columns):
            hdr[j].text = str(c)
        for _, r in sub.iterrows():
            row = table.add_row().cells
            for j, c in enumerate(sub.columns):
                v = r[c]
                if isinstance(v, float):
                    row[j].text = f"{v:.2f}"
                else:
                    row[j].text = str(v)

    fig = RESULTS / "threshold_transfer_pred_vs_obs.png"
    doc.add_heading("Figure: predicted vs observed", level=1)
    if fig.exists():
        doc.add_paragraph(fig.name)
        doc.add_picture(str(fig), width=Inches(6.5))
    else:
        doc.add_paragraph("(threshold_transfer_pred_vs_obs.png not found)")

    out = RESULTS / "Threshold_Transfer_Model_Paradigm_Summary.docx"
    doc.save(out)
    return out


def main() -> None:
    print(build_docx())


if __name__ == "__main__":
    main()
