from __future__ import annotations

import os
from pathlib import Path

import matplotlib
import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt

matplotlib.use("Agg")
import matplotlib.pyplot as plt


REPO_ROOT = Path(__file__).resolve().parents[1]
RESULTS = REPO_ROOT / "results"


def load_data() -> tuple[pd.DataFrame, pd.DataFrame, pd.DataFrame]:
    m = pd.read_csv(RESULTS / "ibm_fisher_thresholds_marrakesh40960.csv")
    t = pd.read_csv(RESULTS / "ibm_fisher_thresholds_torino40960.csv")
    all_norm = pd.read_csv(RESULTS / "ibm_fisher_thresholds_normalized.csv")
    return m, t, all_norm


def add_preformatted(doc: Document, text: str) -> None:
    for line in text.splitlines():
        p = doc.add_paragraph()
        run = p.add_run(line if line else " ")
        run.font.name = "Courier New"
        run.font.size = Pt(8)


def read_text_if_exists(path: Path) -> str | None:
    if not path.exists():
        return None
    return path.read_text(encoding="utf-8", errors="replace")


def add_regime(df: pd.DataFrame) -> pd.DataFrame:
    out = df.copy()
    out["regime_family"] = out["circuit"].map(
        {"ghz": "GHZ", "hadamard": "HadamardLayers", "layers": "HadamardLayers", "identity": "Identity"}
    )
    return out


def make_threshold_compare_plot(m: pd.DataFrame, t: pd.DataFrame) -> Path:
    merged = m.merge(
        t,
        on=["circuit", "n_qubits"],
        suffixes=("_marrakesh", "_torino"),
    ).sort_values(["circuit", "n_qubits"])

    labels = [f"{r.circuit}_{int(r.n_qubits)}q" for r in merged.itertuples()]
    x = np.arange(len(labels))
    w = 0.38

    fig, ax = plt.subplots(figsize=(11, 4.5))
    ax.bar(x - w / 2, merged["threshold_estimate_marrakesh"], w, label="Marrakesh")
    ax.bar(x + w / 2, merged["threshold_estimate_torino"], w, label="Torino")
    ax.set_ylabel("Estimated threshold N*")
    ax.set_title("IBM Fisher threshold per circuit")
    ax.set_xticks(x)
    ax.set_xticklabels(labels, rotation=35, ha="right")
    ax.grid(axis="y", alpha=0.25)
    ax.legend()
    fig.tight_layout()

    out = RESULTS / "ibm_threshold_compare_bar.png"
    fig.savefig(out, dpi=250)
    plt.close(fig)
    return out


def make_regime_boxplots(all_norm: pd.DataFrame) -> tuple[Path, Path]:
    df = add_regime(all_norm)
    order = ["GHZ", "Identity", "HadamardLayers"]
    backends = ["marrakesh", "torino"]

    # Raw N* by regime
    fig, ax = plt.subplots(figsize=(10, 4.5))
    pos = np.arange(len(order))
    width = 0.34
    for i, backend in enumerate(backends):
        data = [df[(df["backend"] == backend) & (df["regime_family"] == reg)]["threshold_estimate"].values for reg in order]
        bp = ax.boxplot(
            data,
            positions=pos + (i - 0.5) * width,
            widths=0.28,
            patch_artist=True,
            manage_ticks=False,
        )
        color = "#1f77b4" if backend == "marrakesh" else "#ff7f0e"
        for box in bp["boxes"]:
            box.set(facecolor=color, alpha=0.5)
    ax.set_xticks(pos)
    ax.set_xticklabels(order)
    ax.set_ylabel("Estimated threshold N*")
    ax.set_title("Threshold distribution by regime family")
    ax.grid(axis="y", alpha=0.25)
    ax.legend(backends, loc="upper right")
    fig.tight_layout()
    out1 = RESULTS / "ibm_threshold_regime_boxplot.png"
    fig.savefig(out1, dpi=250)
    plt.close(fig)

    # N*/qubit by regime
    fig, ax = plt.subplots(figsize=(10, 4.5))
    for i, backend in enumerate(backends):
        data = [df[(df["backend"] == backend) & (df["regime_family"] == reg)]["n_per_qubit"].values for reg in order]
        bp = ax.boxplot(
            data,
            positions=pos + (i - 0.5) * width,
            widths=0.28,
            patch_artist=True,
            manage_ticks=False,
        )
        color = "#1f77b4" if backend == "marrakesh" else "#ff7f0e"
        for box in bp["boxes"]:
            box.set(facecolor=color, alpha=0.5)
    ax.set_xticks(pos)
    ax.set_xticklabels(order)
    ax.set_ylabel("N* / qubit")
    ax.set_title("Normalized threshold by regime family")
    ax.grid(axis="y", alpha=0.25)
    ax.legend(backends, loc="upper right")
    fig.tight_layout()
    out2 = RESULTS / "ibm_threshold_regime_nperqubit_boxplot.png"
    fig.savefig(out2, dpi=250)
    plt.close(fig)
    return out1, out2


def csv_table_from_df(doc: Document, df: pd.DataFrame, title: str) -> None:
    doc.add_heading(title, level=2)
    if df.empty:
        doc.add_paragraph("(empty)")
        return
    cols = list(df.columns)
    table = doc.add_table(rows=1, cols=len(cols))
    hdr = table.rows[0].cells
    for j, c in enumerate(cols):
        hdr[j].text = str(c)
    for _, r in df.iterrows():
        row = table.add_row().cells
        for j, c in enumerate(cols):
            v = r[c]
            if isinstance(v, float):
                row[j].text = f"{v:.4g}" if abs(v) < 1e5 else f"{v:.0f}"
            else:
                row[j].text = str(v)


def build_docx(m: pd.DataFrame, t: pd.DataFrame, all_norm: pd.DataFrame, figures: list[Path]) -> Path:
    doc = Document()
    doc.add_heading("IBM Fisher Threshold Summary", level=0)
    doc.add_paragraph("Dataset: existing IBM 40,960-shot runs on Marrakesh and Torino. Method: Fisher trace sweep over prefix lengths with threshold estimated at max |d(trace)/dN|.")

    # High-level stats
    m_mean = float(m["threshold_estimate"].mean())
    t_mean = float(t["threshold_estimate"].mean())
    m_std = float(m["threshold_estimate"].std())
    t_std = float(t["threshold_estimate"].std())

    doc.add_heading("Key Findings", level=1)
    doc.add_paragraph(f"- Mean threshold is similar across backends: Marrakesh={m_mean:.1f}, Torino={t_mean:.1f}.")
    doc.add_paragraph(f"- Dispersion is high on both backends: Marrakesh std={m_std:.1f}, Torino std={t_std:.1f}.")
    doc.add_paragraph("- GHZ family consistently lands at high thresholds (~36k), while Hadamard/Layers contains low-threshold modes (~1.5k-3k).")
    doc.add_paragraph("- Result supports regime-dependent threshold families more than a single universal raw-N threshold.")

    # Small comparison table
    doc.add_heading("Per-Circuit Thresholds (Marrakesh vs Torino)", level=1)
    cmp_df = m.merge(t, on=["circuit", "n_qubits"], suffixes=("_marrakesh", "_torino")).sort_values(["circuit", "n_qubits"])
    table = doc.add_table(rows=1, cols=5)
    hdr = table.rows[0].cells
    hdr[0].text = "Circuit"
    hdr[1].text = "Qubits"
    hdr[2].text = "Marrakesh N*"
    hdr[3].text = "Torino N*"
    hdr[4].text = "|Delta|"
    for _, r in cmp_df.iterrows():
        row = table.add_row().cells
        row[0].text = str(r["circuit"])
        row[1].text = str(int(r["n_qubits"]))
        row[2].text = f"{float(r['threshold_estimate_marrakesh']):.0f}"
        row[3].text = f"{float(r['threshold_estimate_torino']):.0f}"
        row[4].text = f"{abs(float(r['threshold_estimate_marrakesh']) - float(r['threshold_estimate_torino'])):.0f}"

    backend_path = RESULTS / "ibm_fisher_threshold_backend_summary.csv"
    if backend_path.exists():
        csv_table_from_df(doc, pd.read_csv(backend_path), "Backend aggregates (from sweep CSV)")

    regime_path = RESULTS / "ibm_fisher_thresholds_normalized_regime_summary.csv"
    if regime_path.exists():
        csv_table_from_df(doc, pd.read_csv(regime_path), "Normalized thresholds by regime (mean N*, N*/qubit)")

    m_report = read_text_if_exists(RESULTS / "ibm_fisher_threshold_report_marrakesh40960.txt")
    if m_report:
        doc.add_heading("Appendix A: Marrakesh sweep report (verbatim)", level=1)
        add_preformatted(doc, m_report)

    t_report = read_text_if_exists(RESULTS / "ibm_fisher_threshold_report_torino40960.txt")
    if t_report:
        doc.add_heading("Appendix B: Torino sweep report (verbatim)", level=1)
        add_preformatted(doc, t_report)

    doc.add_heading("Figures", level=1)
    for fig in figures:
        doc.add_paragraph(fig.name)
        doc.add_picture(str(fig), width=Inches(6.5))

    # Add existing Fisher figure if present
    phase_fig = RESULTS / "fisher_phase_transition_summary.png"
    if phase_fig.exists():
        doc.add_paragraph("fisher_phase_transition_summary.png")
        doc.add_picture(str(phase_fig), width=Inches(6.5))

    out = RESULTS / "IBM_Fisher_Threshold_Summary.docx"
    doc.save(out)
    return out


def main() -> None:
    m, t, all_norm = load_data()
    fig1 = make_threshold_compare_plot(m, t)
    fig2, fig3 = make_regime_boxplots(all_norm)
    out = build_docx(m, t, all_norm, [fig1, fig2, fig3])
    print(out)


if __name__ == "__main__":
    main()
